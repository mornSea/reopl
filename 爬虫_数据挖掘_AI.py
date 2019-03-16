//网站目录爬取
  //0:robots.txt的文件，用于告诉爬虫那些文件夹或者哪些文件是网站的拥有者或者管理员不希望被搜索引擎和爬虫浏览的，或者是不希望被非人类的东西查看的。但是不仅仅如此，在这个文件中，有时候还会指明sitemap的位置
User-agent: *
Disallow: /*?*
Disallow: /trackback
Disallow: /wp-*/
Disallow: */comment-page-*
Disallow: /*?replytocom=*
Disallow: */trackback
Disallow: /?random
Disallow: */feed
Disallow: /*.css$
Disallow: /*.js$
Sitemap: http://www.freebuf.com/sitemap.txt

//1:获取当前页面所有的url：
import urllib
from bs4 import BeautifulSoup
import re
 
def get_all_url(url):
   urls = []
   web = urllib.urlopen(url)
soup =BeautifulSoup(web.read())
#通过正则过滤合理的url(针对与freebuf.com来讲)
   tags_a =soup.findAll(name='a',attrs={'href':re.compile("^https?://")})
    try :
       for tag_a in tags_a:
           urls.append(tag_a['href'])
       #return urls
   except:
       pass
   return  urls
 
#得到所有freebuf.com下的url
def get_local_urls(url):
   local_urls = []
   urls = get_all_url(url)
   for _url in urls:
       ret = _url
       if 'freebuf.com' in ret.replace('//','').split('/')[0]:
           local_urls.append(_url)
   return  local_urls
 
#得到所有的不是freebuf.com域名的url
def get_remote_urls(url):
   remote_urls = []
   urls = get_all_url(url)
   for _url in urls:
       ret = _url
       if "freebuf.com" not in ret.replace('//','').split('/')[0]:
           remote_urls.append(_url)
   return  remote_urls
 
def __main__():
   url = 'http://freebuf.com/'
   rurls = get_remote_urls(url)
   print "--------------------remote urls-----------------------"
   for ret in rurls:
       print ret
   print "---------------------localurls-----------------------"   
   lurls = get_local_urls(url)
   for ret in lurls:
       print ret
   
   
if __name__ == '__main__':
__main__()

//2:sitemap爬虫
import urllib
from bs4 import BeautifulSoup
import urlparse
import time
import urllib2
 
url = "http://xxxx.xx/"
domain = "xxxx.xx"
deep = 0
tmp = ""
sites = set()
visited = set()
#local = set()
def get_local_pages(url,domain):
   global deep
   global sites
   global tmp
   repeat_time = 0
   pages = set()
   
    #防止url读取卡住
   while True:
       try:
           print "Ready to Open the web!"
           time.sleep(1)
           print "Opening the web", url
           web = urllib2.urlopen(url=url,timeout=3)
           print "Success to Open the web"
           break
       except:
           print "Open Url Failed !!! Repeat"
           time.sleep(1)
           repeat_time = repeat_time+1
           if repeat_time == 5:
                return
   
   
   print "Readint the web ..."
   soup = BeautifulSoup(web.read())
   print "..."
   tags = soup.findAll(name='a')
   for tag in tags:
       
       #避免参数传递异常
       try:
           ret = tag['href']
       except:
           print "Maybe not the attr : href"
           continue
       o = urlparse.urlparse(ret)
       """
       #Debug I/O
       for _ret in o:
           if _ret == "":
                pass
           else:
                print _ret
       """
       #处理相对路径url
       if o[0] is "" and o[1] is "":
           print "Fix  Page: " +ret
           url_obj = urlparse.urlparse(web.geturl())
           ret = url_obj[0] + "://" + url_obj[1] + url_obj[2] + ret
           #保持url的干净
           ret = ret[:8] + ret[8:].replace('//','/')
           o = urlparse.urlparse(ret)
                     #这里不是太完善，但是可以应付一般情况
           if '../' in o[2]:
                paths = o[2].split('/')
               for i inrange(len(paths)):
                    if paths[i] == '..':
                        paths[i] = ''
                        if paths[i-1]:
                            paths[i-1] = ''
                tmp_path = ''
                for path in paths:
                    if path == '':
                        continue
                    tmp_path = ret_path + '/' +path
                ret =ret.replace(o[2],ret_path)
           print "FixedPage: " + ret
           
           
       #协议处理
       if 'http' not in o[0]:
           print "Bad  Page：" + ret.encode('ascii')
           continue
       
       #url合理性检验
       if o[0] is "" and o[1] is not "":
           print "Bad  Page: " +ret
           continue
       
       #域名检验
       if domain not in o[1]:
           print "Bad  Page: " +ret
           continue
       
       #整理，输出
       newpage = ret
       if newpage not in sites:
           print "Add New Page: " + newpage
           pages.add(newpage)
   return pages
 
#dfs算法遍历全站
def dfs(pages):
    #无法获取新的url说明便利完成，即可结束dfs
   if pages is set():
       return
   global url
   global domain
   global sites
   global visited
   sites = set.union(sites,pages)
   for page in pages:
       if page not in visited:
           print "Visiting",page
           visited.add(page)
           url = page
           pages = get_local_pages(url, domain)
           dfs(pages)
   
   print "sucess"
 
 
pages = get_local_pages(url, domain)
dfs(pages)
for i in sites:
print i

//3:Web元素处理
Soup = BeautifulSoup(data)#构建一个解析器
Tags = Soup.findAll(name,attr)

//定位到图片，那么我们怎么样来寻找与保存
from docx import Document
from bs4 import BeautifulSoup
import urllib
 
url ="http://freebuf.com/news/94263.html"
data = urllib.urlopen(url)
 
document = Document()
 
soup = BeautifulSoup(data)
article = soup.find(name ="div",attrs={'id':'contenttxt'}).children
for e in article:
   try:
       if e.img:
           pic_name = ''
           print e.img.attrs['src']
           if 'gif' in e.img.attrs['src']:
                pic_name = 'temp.gif'
           elif 'png' in e.img.attrs['src']:
                pic_name = 'temp.png'
           elif 'jpg' in e.img.attrs['src']:
                pic_name = 'temp.jpg'
           else:
                pic_name = 'temp.jpeg'
           urllib.urlretrieve(e.img.attrs['src'], filename=pic_name)
           document.add_picture(pic_name)
   except:
       pass
   if e.string:
       print e.string.encode('gbk','ignore')
       document.add_paragraph(e.string)
       
document.save("freebuf_article.docx")
print "success create a document"




//数据挖掘
//1:一，定义：文本挖掘：从大量文本数据中抽取出有价值的知识，并且利用这些知识重新组织信息的过程。    二，语料库（语料库）语料库是我们要分析的所有文档的集合。

import os
import os.path
 
filePaths = []  #定义一个数组变量
#再用OS.walk的方法传入目录
#文件所在的文件目录，命名为root
#root下的所有子目录，命名为dirs
#root下的所有文件，命名为files
for root, dirs, files in os.walk(
    "C:\\Python_DM\\2.1\\SogouC.mini\\Sample"
):
    #进行遍历，需要得到输入目录下的所有文件
    for name in files:  #为了拿到root目录下的所有文件，我们再次便利所有的文件（代码：for name in files:）把它追加到filePaths变量中去即可。
        filePaths.append(os.path.join(root, name)) 
"""
os.path.join,拼接文件路径的方法。如果没有name，则filepaths里面没有xx.txt文件，没有root则没有文件目录路径。
"""
 
import codecs  #编码转换
 
filePaths = [];
fileContents = [];
for root, dirs, files in os.walk(
    "C:\\Python_DM\\2.1\\SogouC.mini\\Sample"
):
    for name in files:
        filePath = os.path.join(root, name);
        filePaths.append(filePath);
        f = codecs.open(filePath, 'r', 'utf-8') #1.文件路径 2.打开方式 3.文件编码
        fileContent = f.read()
        f.close()
        fileContents.append(fileContent)
 
import pandas;
corpos = pandas.DataFrame({
    'filePath': filePaths, 
    'fileContent': fileContents
})


//2:中文分词：将一个汉字序列切分成一个一个单独的词。

import jieba;

for w in jieba.cut("我爱Python"):
    print(w)

//解霸最主要的方法是切方法：  jieba.cut方法接受两个输入参数：cut_all参数用来控制是否采用全模式 及字符串。jieba.cut_for_search方法接受一个参数：需要分词的字符串，该方法适合用于搜索引擎构建倒排索引的分词，粒度比较细    jieba.cut以及jieba.cut_for_search返回的结构都是一个可迭代的generator，可以使用的循环来获得分词后得到的每一个词语（Unicode）的，也可以用列表（jieba.cut（...））转化为列表代码示例（分词）

//用于专业的场景：
import jieba;
seg_list = jieba.cut(
    "真武七截阵和天罡北斗阵哪个更厉害呢？"
)
for w in seg_list:
    print(w)

//为了改善这个现象，我们用导入词库的方法
import jieba;
jieba.add_word('真武七截阵')#添加词库
jieba.add_word('天罡北斗阵')
seg_list = jieba.cut(
    "真武七截阵和天罡北斗阵哪个更厉害呢？"
)
for w in seg_list:
    print(w)

//单词多 低效，我们可以用jieba.load_userdict(‘D:\\PDM\\2.2\\金庸武功招式.txt’)方法一次性导入整个词库，txt文件中为每行一个特定的词。


//3:词频统计
# -*- coding: utf-8 -*-
"""搭建语料库，及分词"""
import os
import os.path
import codecs
 
filePaths = []
fileContents = []
for root, dirs, files in os.walk(
    "C:\\Python_DM\\2.1\\SogouC.mini\\Sample"
):
    for name in files:
        filePath = os.path.join(root, name)
        filePaths.append(filePath);
        f = codecs.open(filePath, 'r', 'utf-8')
        fileContent = f.read()
        f.close()
        fileContents.append(fileContent)
 
import pandas;
corpos = pandas.DataFrame({
    'filePath': filePaths, 
    'fileContent': fileContents
});
 
import jieba
 
segments = []
filePaths = []
for index, row in corpos.iterrows():
    filePath = row['filePath']
    fileContent = row['fileContent']
    segs = jieba.cut(fileContent)
    for seg in segs:
        segments.append(seg)
        filePaths.append(filePath)
 
segmentDataFrame = pandas.DataFrame({
    'segment': segments, 
    'filePath': filePaths
})
 
"""filepath一直没变，对fileContents进行了分词，赋值如segment."""
 
 
import numpy
#进行词频统计        
segStat = segmentDataFrame.groupby( #调用groupby方法
            by="segment"
        )["segment"].agg({
            "计数":numpy.size
        }).reset_index().sort_values(#重新设置索引
            by=['计数'],
            ascending=False#倒序排序
        )
"""排在前面的为停用词"""
 
#移除停用词
stopwords = pandas.read_csv(
    "C:\\Python_DM\\2.3\\StopwordsCN.txt", 
    encoding='utf8', 
    index_col=False
)
 
#获得没有停用词的词频统计结果
fSegStat = segStat[
    ~segStat.segment.isin(stopwords.stopword)
]#'~'取反，不包含停用词的留下。

//移除停用词的另一种方法，加if判断
import os
import os.path
import codecs
 
filePaths = []
fileContents = []
for root, dirs, files in os.walk(
    "C:\\Python_DM\\2.1\\SogouC.mini\\Sample"
):
    for name in files:
        filePath = os.path.join(root, name)
        filePaths.append(filePath);
        f = codecs.open(filePath, 'r', 'utf-8')
        fileContent = f.read()
        f.close()
        fileContents.append(fileContent)
 
import pandas;
corpos = pandas.DataFrame({
    'filePath': filePaths, 
    'fileContent': fileContents
});

import jieba
 
segments = []
filePaths = []
for index, row in corpos.iterrows():
    filePath = row['filePath']
    fileContent = row['fileContent']
    segs = jieba.cut(fileContent)
    for seg in segs:#下面加一个判断，如果不在分词中就加入进分组，加一个条件：分词去除空格后的长度大于0
        if seg not in stopwords.stopword.values and len(seg.strip())>0:
            segments.append(seg)
            filePaths.append(filePath)
 
segmentDataFrame = pandas.DataFrame({
    'segment': segments, 
    'filePath': filePaths
});
 
segStat = segmentDataFrame.groupby(
            by="segment"
        )["segment"].agg({
            "计数":numpy.size
        }).reset_index().sort_values(
            by=["计数"],
            ascending=False
        );


//分组统计：
DataFrame.groupby(
        by=列名数组)[统计列明数组].agg({
                '统计项名称':统计函数
                })

//取反：（对布尔值）
df[~df.列名.isin(数组)]


//3:词云绘制
//安装词云工具包   这个地址：https://www.lfd.uci.edu/~gohlke/pythonlibs/
# -*- coding: utf-8 -*-
import os;
import os.path;
import codecs;
 
filePaths = [];
fileContents = [];
for root, dirs, files in os.walk(
    "C:\\Python_DM\\2.4\\SogouC.mini\\Sample\\C000007"
):
    for name in files:
        filePath = os.path.join(root, name);
        filePaths.append(filePath);
        f = codecs.open(filePath, 'r', 'utf-8')
        fileContent = f.read()
        f.close()
        fileContents.append(fileContent)
 
import pandas;
corpos = pandas.DataFrame({
    'filePath': filePaths, 
    'fileContent': fileContents
});
 
import jieba
 
segments = []
filePaths = []
for index, row in corpos.iterrows():
    filePath = row['filePath']
    fileContent = row['fileContent']
    segs = jieba.cut(fileContent)
    for seg in segs:
        segments.append(seg)
        filePaths.append(filePath)
 
segmentDataFrame = pandas.DataFrame({
    'segment': segments, 
    'filePath': filePaths
});
 
 
import numpy;
#进行词频统计        
segStat = segmentDataFrame.groupby(
            by="segment"
        )["segment"].agg({
            "计数":numpy.size
        }).reset_index().sort_values(
            by=["计数"],
            ascending=False
        );
 
#移除停用词
stopwords = pandas.read_csv(
    "C:\\Python_DM\\2.4\\StopwordsCN.txt", 
    encoding='utf8', 
    index_col=False
)
 
fSegStat = segStat[
    ~segStat.segment.isin(stopwords.stopword)
]
 
#绘画词云
#http://www.lfd.uci.edu/~gohlke/pythonlibs/
from wordcloud import WordCloud
import matplotlib.pyplot as plt
 
#传入字体文件的路径，还有背景颜色
#微软雅黑的字体，黑色背景
wordcloud = WordCloud(
    font_path='C:\\Python_DM\\2.4\\simhei.ttf', 
    background_color="black"
)
 
#把分词设置成数据框的索引，再调用to_dict()的方法，获得一个字典的数据结构了。
words = fSegStat.set_index('segment').to_dict()
 
wordcloud.fit_words(words['计数'])
 
plt.imshow(wordcloud)
 
#plt.close()

//美化词云（词云放入某图片形象中）



 
//5:关键词提取
import os
import codecs
import pandas
import jieba
import jieba.analyse #关键字提取方法
 
#定义好存储的列数组，抽取5个关键字
filePaths = []
contents = []
tag1s = []
tag2s = []
tag3s = []
tag4s = []
tag5s = []
 
for root, dirs, files in os.walk(
    "C:\\Python_DM\\2.6\\SogouC.mini\\Sample\\"
):
    for name in files:
        filePath = root + '\\' + name;
        f = codecs.open(filePath, 'r', 'utf-8')
        content = f.read().strip()
        f.close()
        
        #将文件内容（content）传递给extract_tags方法
        #把管家你在tags添加到对应的列中
        tags = jieba.analyse.extract_tags(content, topK=5)
        filePaths.append(filePath)
        contents.append(content)
        tag1s.append(tags[0])
        tag2s.append(tags[1])
        tag3s.append(tags[2])
        tag4s.append(tags[3])
        tag5s.append(tags[4])
 
tagDF = pandas.DataFrame({
    'filePath': filePaths, 
    'content': contents, 
    'tag1': tag1s, 
    'tag2': tag2s, 
    'tag3': tag3s, 
    'tag4': tag4s, 
    'tag5': tag5s
})

//计算公式： TF = 该次在文档中出现的次数  IDF就是每个词的权重，它的大小与一个词的常见程度成反比     TF - IDF = TF * IDF


//数据挖掘：垃圾信息识别
//0:准备
%matplotlib inline

import matplotlib.pyplot as plt
import csv
from textblob import TextBlob
import pandas
import sklearn
import cPickle
import numpy as np
from sklearn.feature_extraction.text import CountVectorizer, TfidfTransformer
from sklearn.naive_bayes import MultinomialNB
from sklearn.svm import SVC, LinearSVC
from sklearn.metrics import classification_report, f1_score, accuracy_score, confusion_matrix
from sklearn.pipeline import Pipeline
from sklearn.grid_search import GridSearchCV
from sklearn.cross_validation import StratifiedKFold, cross_val_score, train_test_split 
from sklearn.tree import DecisionTreeClassifier 
from sklearn.learning_curve import learning_curve

//打印 SMS 语料库中的前 10 条信息：
for message_no, message in enumerate(messages[:10]):
    print message_no, message

//Pandas 库替我们处理 TSV 或CSV / Excel文件:
messages = pandas.read_csv('./data/SMSSpamCollection', sep='t', quoting=csv.QUOTE_NONE,
                           names=["label", "message"])

//2:用 pandas 轻松查看统计信息：信息的长度是多少：哪些是超长信息？

//写一个将信息分割成单词的函数：
def split_into_tokens(message):
    message = unicode(message, 'utf8')  # convert bytes into proper unicode
    return TextBlob(message).words
    

//自然语言处理（NLP）的问题：大写字母是否携带信息？单词的不同形式（“goes”和“go”）是否携带信息？叹词和限定词是否携带信息？换句话说，我们想对文本进行更好的标准化。    我们使用 textblob 获取 part-of-speech (POS) 标签：
TextBlob("Hello world, how is it going?").tags  # list of (word, POS) pairs
1
TextBlob("Hello world, how is it going?").tags 

//并将单词标准化为基本形式 (lemmas)：
def split_into_lemmas(message):
    message = unicode(message, 'utf8').lower()
    words = TextBlob(message).words
    # for each word, take its "base form" = lemma 
    return [word.lemma for word in words]
 
messages.message.head().apply(split_into_lemmas)


//3:数据转换为向量:
//用词袋模型完成这项工作需要三个步骤：    1.  对每个词在每条信息中出现的次数进行计数（词频）；    2. 对计数进行加权，这样经常出现的单词将会获得较低的权重（逆向文件频率）；    3. 将向量由原始文本长度归一化到单位长度（L2 范式）。        每个向量的维度等于 SMS 语料库中包含的独立词的数量。

//用scikit-learn (sklearn)
bow_transformer = CountVectorizer(analyzer=split_into_lemmas).fit(messages['message'])
print len(bow_transformer.vocabulary_)

bow4 = bow_transformer.transform([message4])
print bow4
print bow4.shape


//可用性检测，哪些词出现了多次？
print bow_transformer.get_feature_names()[6736]
print bow_transformer.get_feature_names()[8013]

//整个 SMS 语料库的词袋计数是一个庞大的稀疏矩阵：
messages_bow = bow_transformer.transform(messages['message'])
print 'sparse matrix shape:', messages_bow.shape

print 'number of non-zeros:', messages_bow.nnz

print 'sparsity: %.2f%%' % (100.0 * messages_bow.nnz / (messages_bow.shape[0] * messages_bow.shape[1]))

//最终，计数后，使用 scikit-learn 的 TFidfTransformer 实现的 TF-IDF 完成词语加权和归一化。
tfidf_transformer = TfidfTransformer().fit(messages_bow)
tfidf4 = tfidf_transformer.transform(bow4)
print tfidf4

print tfidf_transformer.idf_[bow_transformer.vocabulary_['u']]
print tfidf_transformer.idf_[bow_transformer.vocabulary_['university']]


//4:训练模型,检测垃圾信息
//用 scikit-learn，首先选择 Naive Bayes 分类器：
%time spam_detector = MultinomialNB().fit(messages_tfidf, messages['label'])

CPU times: user 4.51 ms, sys: 987 µs, total: 5.49 ms

Wall time: 4.77 ms

//分类一个随机信息：
print 'predicted:', spam_detector.predict(tfidf4)[0]
print 'expected:', messages.label[3]


//训练的准确率将非常接近 100%，但是我们不能用它来分类任何新信息.一个正确的做法是将数据分为训练集和测试集，在模型拟合和调参时只能使用训练数据，不能以任何方式使用测试数据，通过这个方法确保模型没有“作弊”，最终使用测试数据评价模型可以代表模型真正的预测性能。

//
msg_train, msg_test, label_train, label_test = 
    train_test_split(messages['message'], messages['label'], test_size=0.2)
 
print len(msg_train), len(msg_test), len(msg_train) + len(msg_test)


//回顾整个流程，将所有步骤放入 scikit-learn 的 Pipeline 中:
def split_into_lemmas(message):
    message = unicode(message, 'utf8').lower()
    words = TextBlob(message).words
    # for each word, take its "base form" = lemma 
    return [word.lemma for word in words]
 
pipeline = Pipeline([
    ('bow', CountVectorizer(analyzer=split_into_lemmas)),  # strings to token integer counts
    ('tfidf', TfidfTransformer()),  # integer counts to weighted TF-IDF scores
    ('classifier', MultinomialNB()),  # train on TF-IDF vectors w/ Naive Bayes classifier
])


//Naive Bayes 是一个高偏差-低方差的分类器（简单且稳定，不易过度拟合）用它：
def plot_learning_curve(estimator, title, X, y, ylim=None, cv=None,
                        n_jobs=-1, train_sizes=np.linspace(.1, 1.0, 5)):
    plt.figure()
    plt.title(title)
    if ylim is not None:
        plt.ylim(*ylim)
    plt.xlabel("Training examples")
    plt.ylabel("Score")
    train_sizes, train_scores, test_scores = learning_curve(
        estimator, X, y, cv=cv, n_jobs=n_jobs, train_sizes=train_sizes)
    train_scores_mean = np.mean(train_scores, axis=1)
    train_scores_std = np.std(train_scores, axis=1)
    test_scores_mean = np.mean(test_scores, axis=1)
    test_scores_std = np.std(test_scores, axis=1)
    plt.grid()
 
    plt.fill_between(train_sizes, train_scores_mean - train_scores_std,
                     train_scores_mean + train_scores_std, alpha=0.1,
                     color="r")
    plt.fill_between(train_sizes, test_scores_mean - test_scores_std,
                     test_scores_mean + test_scores_std, alpha=0.1, color="g")
    plt.plot(train_sizes, train_scores_mean, 'o-', color="r",
             label="Training score")
    plt.plot(train_sizes, test_scores_mean, 'o-', color="g",
             label="Cross-validation score")
 
    plt.legend(loc="best")
    return plt


%time plot_learning_curve(pipeline, "accuracy vs. training set size", msg_train, label_train, cv=5)

%time plot_learning_curve(pipeline, "accuracy vs. training set size", msg_train, label_train, cv=5)
Python

CPU times: user 382 ms, sys: 83.1 ms, total: 465 ms
Wall time: 28.5 s

CPU times: user 382 ms, sys: 83.1 ms, total: 465 ms
Wall time: 28.5 s

&lt;module 'matplotlib.pyplot' from '/Volumes/work/workspace/vew/sklearn_intro/lib/python2.7/site-packages/matplotlib/pyplot.pyc'&gt;

&lt;module 'matplotlib.pyplot' from '/Volumes/work/workspace/vew/sklearn_intro/lib/python2.7/site-packages/matplotlib/pyplot.pyc'&gt;


//5:如何调整参数？
//我们会问：IDF 加权对准确性有什么影响？消耗额外成本进行词形还原（与只用纯文字相比）真的会有效果吗？
params = {
    'tfidf__use_idf': (True, False),
    'bow__analyzer': (split_into_lemmas, split_into_tokens),
}
 
grid = GridSearchCV(
    pipeline,  # pipeline from above
    params,  # parameters to tune via cross validation
    refit=True,  # fit using all available data at the end, on the best found param combination
    n_jobs=-1,  # number of cores to use for parallelization; -1 for "all cores"
    scoring='accuracy',  # what score are we optimizing?
    cv=StratifiedKFold(label_train, n_folds=5),  # what type of cross validation to use
)


%time nb_detector = grid.fit(msg_train, label_train)
 
print nb_detector.grid_scores_




CPU times: user 4.09 s, sys: 291 ms, total: 4.38 s
Wall time: 20.2 s
[mean: 0.94752, std: 0.00357, params: {'tfidf__use_idf': True, 'bow__analyzer': &lt;function split_into_lemmas at 0x1131e8668&gt;}, mean: 0.92958, std: 0.00390, params: {'tfidf__use_idf': False, 'bow__analyzer': &lt;function split_into_lemmas at 0x1131e8668&gt;}, mean: 0.94528, std: 0.00259, params: {'tfidf__use_idf': True, 'bow__analyzer': &lt;function split_into_tokens at 0x11270b7d0&gt;}, mean: 0.92868, std: 0.00240, params: {'tfidf__use_idf': False, 'bow__analyzer': &lt;function split_into_tokens at 0x11270b7d0&gt;}]

//（首先显示最佳参数组合：在这个案例中是使用 idf=True 和 analyzer=split_into_lemmas 的参数组合）    快速合理性检查:
print nb_detector.predict_proba(["Hi mom, how are you?"])[0]
print nb_detector.predict_proba(["WINNER! Credit for free!"])[0]


//让我们尝试另一个分类器：支持向量机（SVM）。SVM 可以非常迅速的得到结果，它所需要的参数调整也很少  尤其在文本处理表现出色:
pipeline_svm = Pipeline([
    ('bow', CountVectorizer(analyzer=split_into_lemmas)),
    ('tfidf', TfidfTransformer()),
    ('classifier', SVC()),  # &lt;== change here
])
 
# pipeline parameters to automatically explore and tune
param_svm = [
  {'classifier__C': [1, 10, 100, 1000], 'classifier__kernel': ['linear']},
  {'classifier__C': [1, 10, 100, 1000], 'classifier__gamma': [0.001, 0.0001], 'classifier__kernel': ['rbf']},
]
 
grid_svm = GridSearchCV(
    pipeline_svm,  # pipeline from above
    param_grid=param_svm,  # parameters to tune via cross validation
    refit=True,  # fit using all data, on the best detected classifier
    n_jobs=-1,  # number of cores to use for parallelization; -1 for "all cores"
    scoring='accuracy',  # what score are we optimizing?
    cv=StratifiedKFold(label_train, n_folds=5),  # what type of cross validation to use
)

%time svm_detector = grid_svm.fit(msg_train, label_train) # find the best combination from param_svm
 
print svm_detector.grid_scores_

CPU times: user 5.24 s, sys: 170 ms, total: 5.41 s
Wall time: 1min 8s
[mean: 0.98677, std: 0.00259, params: {'classifier__kernel': 'linear', 'classifier__C': 1}, mean: 0.98654, std: 0.00100, params: {'classifier__kernel': 'linear', 'classifier__C': 10}, mean: 0.98654, std: 0.00100, params: {'classifier__kernel': 'linear', 'classifier__C': 100}, mean: 0.98654, std: 0.00100, params: {'classifier__kernel': 'linear', 'classifier__C': 1000}, mean: 0.86432, std: 0.00006, params: {'classifier__gamma': 0.001, 'classifier__kernel': 'rbf', 'classifier__C': 1}, mean: 0.86432, std: 0.00006, params: {'classifier__gamma': 0.0001, 'classifier__kernel': 'rbf', 'classifier__C': 1}, mean: 0.86432, std: 0.00006, params: {'classifier__gamma': 0.001, 'classifier__kernel': 'rbf', 'classifier__C': 10}, mean: 0.86432, std: 0.00006, params: {'classifier__gamma': 0.0001, 'classifier__kernel': 'rbf', 'classifier__C': 10}, mean: 0.97040, std: 0.00587, params: {'classifier__gamma': 0.001, 'classifier__kernel': 'rbf', 'classifier__C': 100}, mean: 0.86432, std: 0.00006, params: {'classifier__gamma': 0.0001, 'classifier__kernel': 'rbf', 'classifier__C': 100}, mean: 0.98722, std: 0.00280, params: {'classifier__gamma': 0.001, 'classifier__kernel': 'rbf', 'classifier__C': 1000}, mean: 0.97040, std: 0.00587, params: {'classifier__gamma': 0.0001, 'classifier__kernel': 'rbf', 'classifier__C': 1000}]

//再一次合理性检查：
print svm_detector.predict(["Hi mom, how are you?"])[0]
print svm_detector.predict(["WINNER! Credit for free!"])[0]


//6:生成预测器     最终的预测器可以序列化到磁盘，以便我们下次想使用它时，可以跳过所有训练直接使用训练好的模型：

# store the spam detector to disk after training
with open('sms_spam_detector.pkl', 'wb') as fout:
    cPickle.dump(svm_detector, fout)
 
# ...and load it back, whenever needed, possibly on a different machine
svm_detector_reloaded = cPickle.load(open('sms_spam_detector.pkl'))

//加载的结果是一个与原始对象表现相同的对象：
print 'before:', svm_detector.predict([message4])[0]
print 'after:', svm_detector_reloaded.predict([message4])[0]





//-----------------------------数据挖掘 与 AI    之间并不清楚的分割线-----------------------------------//




//分析弹幕/评论
import requests
from fake_useragent import UserAgent
import json
import time
import pandas as pd

#下载网页评论数据
def get_page_json(url):
    try:
        ua = UserAgent(verify_ssl=False)
        headers = {"User-Agent": ua.random}
        json_comment = requests.get(url,headers=headers).text
        return json_comment
    except:
        return None

#解析网页评论数据
def parse_page_json(json_comment):
   try:
       comments = json.loads(json_comment)
   except:
       return "error"

comments_list = []
   #获取当页数据有多少条评论（一般情况下为20条）
   num = len(comments['data']['replies'])

for i in range(num):
       comment = comments['data']['replies'][i]
       comment_list = []
       floor = comment['floor']
       ctime = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(comment['ctime']))#时间转换
       likes = comment['like']
       author = comment['member']['uname']
       sex = comment['member']['sex']
       level = comment['member']['level_info']['current_level']
       content = comment['content']['message'].replace('\n','')#将评论内容中的换行符去掉
       #print(content)
       rcount = comment['rcount']
       comment_list.append(floor)
       comment_list.append(ctime)
       comment_list.append(likes)
       comment_list.append(author)
       comment_list.append(sex)
       comment_list.append(level)
       comment_list.append(content)
       comment_list.append(rcount)

comments_list.append(comment_list)

save_to_csv(comments_list)


def save_to_csv(comments_list):
    data = pd.DataFrame(comments_list)
    #注意存储文件的编码为utf_8_sig，不然会乱码，后期会单独深入讲讲为何为这样（如果为utf-8）
    data.to_csv('春晚鬼畜_1.csv', mode='a', index=False, sep=',', header=False,encoding='utf_8_sig')


def main():
    base_url = "https://api.bilibili.com/x/v2/reply?&type=1&oid=19390801&pn=1"
    #通过首页获取评论总页数
    pages = int(json.loads(get_page_json(base_url))['data']['page']['count'])//20
    for page in range(pages):
        url = "https://api.bilibili.com/x/v2/reply?&type=1&oid=19390801&pn="+str(page)
        json_comment = get_page_json(url)
        parse_page_json(json_comment)
        print("正在保存第%d页" % int(page+1))

if page%20 == 0:
            time.sleep(5)

main()


//神经网络
下面是如何在Python项目中创建神经网络的完整代码：

import numpy as np

class NeuralNetwork():

    def __init__(self):

        # seeding for random number generation

        np.random.seed(1)

        #converting weights to a 3 by 1 matrix with values from -1 to 1 and mean of 0

        self.synaptic_weights = 2 * np.random.random((3, 1)) - 1

    def sigmoid(self, x):

        #applying the sigmoid function

        return 1 / (1 + np.exp(-x))

    def sigmoid_derivative(self, x):

        #computing derivative to the Sigmoid function

        return x * (1 - x)

    def train(self, training_inputs, training_outputs, training_iterations):

        #training the model to make accurate predictions while adjusting weights continually

        for iteration in range(training_iterations):

            #siphon the training data via  the neuron

            output = self.think(training_inputs)

            #computing error rate for back-propagation

            error = training_outputs - output

            #performing weight adjustments

            adjustments = np.dot(training_inputs.T, error * self.sigmoid_derivative(output))

            self.synaptic_weights += adjustments

    def think(self, inputs):

        #passing the inputs via the neuron to get output   

        #converting values to floats

        inputs = inputs.astype(float)

        output = self.sigmoid(np.dot(inputs, self.synaptic_weights))

        return output

if __name__ == "__main__":

    #initializing the neuron class

    neural_network = NeuralNetwork()

    print("Beginning Randomly Generated Weights: ")

    print(neural_network.synaptic_weights)

    #training data consisting of 4 examples--3 input values and 1 output

    training_inputs = np.array([[0,0,1],

                                [1,1,1],

                                [1,0,1],

                                [0,1,1]])

    training_outputs = np.array([[0,1,1,0]]).T

    #training taking place

    neural_network.train(training_inputs, training_outputs, 15000)

    print("Ending Weights After Training: ")

    print(neural_network.synaptic_weights)

    user_input_one = str(input("User Input One: "))

    user_input_two = str(input("User Input Two: "))

    user_input_three = str(input("User Input Three: "))

    print("Considering New Situation: ", user_input_one, user_input_two, user_input_three)

    print("New Output data: ")

    print(neural_network.think(np.array([user_input_one, user_input_two, user_input_three])))

print("Wow, we did it!")


//物体检测

//步骤1：使用python 3.6版搭建Anaconda环境。

conda create -nretinanet python=3.6 anaconda

//步骤2：激活环境并安装必要的软件包。

source activateretinanet

conda installtensorflow numpy scipy opencv pillow matplotlib h5py keras

//步骤3：安装ImageAI库

pip install https://github.com/OlafenwaMoses /ImageAI/releases/download/2.0.1/imageai-2.0.1-py3-none-any.whl

//步骤4：下载获得预测所需的预训练模型。该模型基于RetinaNet（未来文章的主题）。访问链接下载——RetinaNet预训练模型:

https://github.com/OlafenwaMoses/ImageAI/releases/download/1.0/resnet50_coco_best_v2.0.1.h5

//步骤5：将下载的文件复制到当前工作文件夹。

//步骤6：从以下链接下载图像，将图像命名为image.png：

https://s3-ap-south-1.amazonaws.com/av-blog-media/wp-content/uploads/2018/06/I1_2009_09_08_drive_0012_001351-768x223.png

//步骤7：打开jupyter笔记本（在终端中键入jupyter notebook）并运行以下代码：

from imageai.Detection import ObjectDetection

import os

execution_path = os.getcwd()

detector = ObjectDetection()

detector.setModelTypeAsRetinaNet()

detector.setModelPath( os.path.join(execution_path ,"resnet50_coco_best_v2.0.1.h5"))

detector.loadModel()

custom_objects = detector.CustomObjects(person=True, car=False)

detections =detector.detectCustomObjectsFromImage(input_image=os.path.join(execution_path ,"image.png"), output_image_path=os.path.join(execution_path ,"image_new.png"), custom_objects=custom_objects,minimum_percentage_probability=65)

for eachObject in detections:

  print(eachObject["name"] + " : " +eachObject["percentage_probability"] )

  print("--------------------------------")

这将创建一个名为image_new.png的修改后的图像文件，其中包含物体的边界框。

//步骤8：如需要打印图像，请使用以下代码：

from IPython.display import Image

Image("image_new.png")


//预测股票
  //首先让我们加载数据集，定义问题的目标变量:

#import packages

import pandas as pd

import numpy as np

#to plot within notebook

import matplotlib.pyplot as plt

%matplotlib inline

#setting figure size

from matplotlib.pylab import rcParams

rcParams['figure.figsize'] = 20,10

#for normalizing data

from sklearn.preprocessing import MinMaxScaler

scaler = MinMaxScaler(feature_range=(0, 1))

#read the file

df = pd.read_csv('NSE-TATAGLOBAL(1).csv')

#print the head

df.head()

//让我们画出目标变量来理解它在我们的数据集中的分布:

#setting index as date

df['Date'] = pd.to_datetime(df.Date,format='%Y-%m-%d')

df.index = df['Date']

#plot

plt.figure(figsize=(16,8))

plt.plot(df['Close'], label='Close Price history')

//移动平均法：第一步是创建一个只包含日期和收盘价列的数据框，然后将其拆分为训练集和验证集来验证我们的预测。


#creating dataframe with date and the target variable

data = df.sort_index(ascending=True, axis=0)

new_data = pd.DataFrame(index=range(0,len(df)),columns=['Date', 'Close'])

for i in range(0,len(data)):

     new_data['Date'][i] = data['Date'][i]

     new_data['Close'][i] = data['Close'][i]

//在将数据分割为训练和验证时，我们不能使用随机分割，因为这会破坏时间顺序。所以这里我将去年的数据作为验证集，将之前4年的数据作为训练集。

#splitting into train and validation

train = new_data[:987]

valid = new_data[987:]

new_data.shape, train.shape, valid.shape

((1235, 2), (987, 2), (248, 2))

train['Date'].min(), train['Date'].max(), valid['Date'].min(), valid['Date'].max()

(Timestamp('2013-10-08 00:00:00'),

Timestamp('2017-10-06 00:00:00'),

Timestamp('2017-10-09 00:00:00'),

Timestamp('2018-10-08 00:00:00'))

//下一步是为验证集创建预测值，并使用真实值来检查RMSE误差。

#make predictions

preds = []

for i in range(0,248):

    a = train['Close'][len(train)-248+i:].sum() + sum(preds)

    b = a/248

    preds.append(b)
#calculate rmse

rms=np.sqrt(np.mean(np.power((np.array(valid['Close'])-preds),2)))

rms

104.51415465984348

//仅仅检查RMSE并不能帮助我们评估模型预测效果的。让我们通过做图得到更直观的理解。下面是预测值和实际值的曲线图。

#plot

valid['Predictions'] = 0

valid['Predictions'] = preds

plt.plot(train['Close'])

plt.plot(valid[['Close', 'Predictions']])

//线性回归
我们将首先按升序对数据集进行排序，然后创建一个单独的数据集，这样创建的任何新特性都不会影响原始数据。

#setting index as date values

df['Date'] = pd.to_datetime(df.Date,format='%Y-%m-%d')

df.index = df['Date']

#sorting

data = df.sort_index(ascending=True, axis=0)

#creating a separate dataset

new_data = pd.DataFrame(index=range(0,len(df)),columns=['Date', 'Close'])

for i in range(0,len(data)):

    new_data['Date'][i] = data['Date'][i]

    new_data['Close'][i] = data['Close'][i]

#create features

from fastai.structured import  add_datepart

add_datepart(new_data, 'Date')

new_data.drop('Elapsed', axis=1, inplace=True)  #elapsed will be the time stamp

//将把数据分成训练集和验证集来检查模型的性能。

#split into train and validation

train = new_data[:987]

valid = new_data[987:]

x_train = train.drop('Close', axis=1)

y_train = train['Close']

x_valid = valid.drop('Close', axis=1)

y_valid = valid['Close']

#implement linear regression

from sklearn.linear_model import LinearRegression

model = LinearRegression()

model.fit(x_train,y_train)

//显示线性回归结果
#plot

valid['Predictions'] = 0

valid['Predictions'] = preds

valid.index = new_data[987:].index

train.index = new_data[:987].index

plt.plot(train['Close'])

plt.plot(valid[['Close', 'Predictions']])

//k近邻
#importing libraries

from sklearn import neighbors

from sklearn.model_selection import GridSearchCV

from sklearn.preprocessing import MinMaxScaler

scaler = MinMaxScaler(feature_range=(0, 1))

使用上一节中相同的训练和验证集:

#scaling data

x_train_scaled = scaler.fit_transform(x_train)

x_train = pd.DataFrame(x_train_scaled)

x_valid_scaled = scaler.fit_transform(x_valid)

x_valid = pd.DataFrame(x_valid_scaled)

#using gridsearch to find the best parameter

params = {'n_neighbors':[2,3,4,5,6,7,8,9]}

knn = neighbors.KNeighborsRegressor()

model = GridSearchCV(knn, params, cv=5)

#fit the model and make predictions

model.fit(x_train,y_train)

preds = model.predict(x_valid)

#plot

valid['Predictions'] = 0

valid['Predictions'] = preds

plt.plot(valid[['Close', 'Predictions']])

plt.plot(train['Close'])



//利用自动ARIMA建立高性能时间序列模型


from pyramid.arima import auto_arima

data = df.sort_index(ascending=True, axis=0)

train = data[:987]

valid = data[987:]

training = train['Close']

validation = valid['Close']

model = auto_arima(training, start_p=1, start_q=1,max_p=3, max_q=3, m=12,start_P=0, seasonal=True,d=1, D=1, trace=True,error_action='ignore',suppress_warnings=True)

model.fit(training)

forecast = model.predict(n_periods=248)

forecast = pd.DataFrame(forecast,index = valid.index,columns=['Prediction'])

// 使用Facebook的Prophet生成快速准确的时间序列预测

实现

#importing prophet

from fbprophet import Prophet

#creating dataframe

new_data = pd.DataFrame(index=range(0,len(df)),columns=['Date', 'Close'])

for i in range(0,len(data)):

    new_data['Date'][i] = data['Date'][i]

    new_data['Close'][i] = data['Close'][i]

new_data['Date'] = pd.to_datetime(new_data.Date,format='%Y-%m-%d')

new_data.index = new_data['Date']

#preparing data

new_data.rename(columns={'Close': 'y', 'Date': 'ds'}, inplace=True)

#train and validation

train = new_data[:987]

valid = new_data[987:]

#fit the model

model = Prophet()

model.fit(train)

#predictions

close_prices = model.make_future_dataframe(periods=len(valid))

forecast = model.predict(close_prices)


//让我们将LSTM实现为一个黑盒，并检查它在特定数据上的性能。



#importing required libraries

from sklearn.preprocessing import MinMaxScaler

from keras.models import Sequential

from keras.layers import Dense, Dropout, LSTM

#creating dataframe

data = df.sort_index(ascending=True, axis=0)

new_data = pd.DataFrame(index=range(0,len(df)),columns=['Date', 'Close'])

for i in range(0,len(data)):

    new_data['Date'][i] = data['Date'][i]

    new_data['Close'][i] = data['Close'][i]

#setting index

new_data.index = new_data.Date

new_data.drop('Date', axis=1, inplace=True)

#creating train and test sets

dataset = new_data.values

train = dataset[0:987,:]

valid = dataset[987:,:]

#converting dataset into x_train and y_train

scaler = MinMaxScaler(feature_range=(0, 1))

scaled_data = scaler.fit_transform(dataset)

x_train, y_train = [], []

for i in range(60,len(train)):

    x_train.append(scaled_data[i-60:i,0])

    y_train.append(scaled_data[i,0])

x_train, y_train = np.array(x_train), np.array(y_train)

x_train = np.reshape(x_train, (x_train.shape[0],x_train.shape[1],1))

# create and fit the LSTM network

model = Sequential()

model.add(LSTM(units=50, return_sequences=True, input_shape=(x_train.shape[1],1)))

model.add(LSTM(units=50))

model.add(Dense(1))

model.compile(loss='mean_squared_error', optimizer='adam')

model.fit(x_train, y_train, epochs=1, batch_size=1, verbose=2)

#predicting 246 values, using past 60 from the train data

inputs = new_data[len(new_data) - len(valid) - 60:].values

inputs = inputs.reshape(-1,1)

inputs  = scaler.transform(inputs)

X_test = []

for i in range(60,inputs.shape[0]):

    X_test.append(inputs[i-60:i,0])

X_test = np.array(X_test)

X_test = np.reshape(X_test, (X_test.shape[0],X_test.shape[1],1))

closing_price = model.predict(X_test)

closing_price = scaler.inverse_transform(closing_price)






















